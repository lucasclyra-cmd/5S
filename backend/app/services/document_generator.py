import os
import subprocess
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.version import DocumentVersion


def generate_docx(
    structured_content: dict,
    template_config: Optional[dict],
    output_path: str,
) -> str:
    """
    Generate a formatted .docx file from structured content.

    structured_content should have:
    - sections: list of {title, content, level}
    - metadata: dict with optional title, author, date
    """
    doc = DocxDocument()

    # Apply template config defaults
    config = template_config or {}
    font_name = config.get("font", "Arial")
    font_size = config.get("font_size", 11)
    margin_cm = config.get("margin_cm", 2.54)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)

    # Add header if configured
    header_text = config.get("header_text")
    if header_text:
        header = doc.sections[0].header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.size = Pt(9)
            run.font.name = font_name

    # Add footer if configured
    footer_text = config.get("footer_text")
    if footer_text:
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.name = font_name

    # Add metadata title
    metadata = structured_content.get("metadata", {})
    title = metadata.get("title", "")
    if title:
        title_para = doc.add_heading(title, level=0)
        for run in title_para.runs:
            run.font.name = font_name

    # Add sections
    sections = structured_content.get("sections", [])
    for section_data in sections:
        section_title = section_data.get("title", "")
        section_content = section_data.get("content", "")
        level = section_data.get("level", 1)

        # Add section heading
        if section_title:
            heading = doc.add_heading(section_title, level=min(level, 4))
            for run in heading.runs:
                run.font.name = font_name

        # Add section content
        if section_content:
            paragraphs = section_content.split("\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    para = doc.add_paragraph(para_text)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in para.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc.save(output_path)
    return output_path


def convert_to_pdf(docx_path: str, output_path: str) -> str:
    """Convert a .docx file to PDF using LibreOffice headless."""
    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)

    result = subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_path,
        ],
        capture_output=True,
        text=True,
        timeout=60,
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed: {result.stderr}"
        )

    # LibreOffice outputs with the same base name but .pdf extension in output_dir
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")

    # Rename to desired output path if different
    if generated_pdf != output_path:
        os.rename(generated_pdf, output_path)

    return output_path


async def format_document(
    version: DocumentVersion,
    structured_content: dict,
    template_config: Optional[dict],
    storage_path: str,
) -> tuple[str, str]:
    """
    Generate formatted .docx and .pdf files for a document version.
    Returns (docx_path, pdf_path).
    """
    formatted_dir = os.path.join(storage_path, "formatted")
    os.makedirs(formatted_dir, exist_ok=True)

    base_name = f"doc_{version.document_id}_v{version.version_number}"
    docx_path = os.path.join(formatted_dir, f"{base_name}.docx")
    pdf_path = os.path.join(formatted_dir, f"{base_name}.pdf")

    # Generate .docx
    generate_docx(structured_content, template_config, docx_path)

    # Convert to PDF
    try:
        convert_to_pdf(docx_path, pdf_path)
    except Exception:
        # PDF conversion might fail if LibreOffice is not available
        pdf_path = ""

    return docx_path, pdf_path
