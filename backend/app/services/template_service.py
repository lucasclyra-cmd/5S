"""
Template Service — loads .docx templates, injects document content,
populates headers, revision history and approval tables.
"""

import os
import re
import copy
import shutil
import subprocess
import unicodedata
from typing import Optional
from io import BytesIO
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import settings


def _strip_accents(text: str) -> str:
    """Remove all diacritical marks from a string (ã→a, ç→c, é→e, etc.)."""
    nfkd = unicodedata.normalize("NFKD", text)
    return "".join(c for c in nfkd if not unicodedata.category(c).startswith("M"))


# ──────────────────────────────────────────────────────────────
# Image extraction from source .docx
# ──────────────────────────────────────────────────────────────

def extract_images_from_docx(docx_path: str) -> list[dict]:
    """
    Extract all images from a .docx file.
    Returns list of {"filename": ..., "content_type": ..., "data": bytes}
    """
    images = []
    try:
        with ZipFile(docx_path, "r") as zf:
            for name in zf.namelist():
                if name.startswith("word/media/"):
                    data = zf.read(name)
                    ext = os.path.splitext(name)[1].lower()
                    content_type = {
                        ".png": "image/png",
                        ".jpg": "image/jpeg",
                        ".jpeg": "image/jpeg",
                        ".gif": "image/gif",
                        ".bmp": "image/bmp",
                        ".tiff": "image/tiff",
                    }.get(ext, "application/octet-stream")
                    images.append({
                        "filename": os.path.basename(name),
                        "content_type": content_type,
                        "data": data,
                    })
    except Exception:
        pass
    return images


# ──────────────────────────────────────────────────────────────
# Placeholder replacement in document
# ──────────────────────────────────────────────────────────────

PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


def _replace_in_paragraph(paragraph, replacements: dict[str, str]):
    """Replace {{PLACEHOLDER}} in a paragraph's runs, preserving formatting."""
    full_text = paragraph.text
    if "{{" not in full_text:
        return False

    replaced = False
    for key, value in replacements.items():
        marker = "{{" + key + "}}"
        if marker in full_text:
            full_text = full_text.replace(marker, value or "")
            replaced = True

    if replaced:
        # Preserve first run's formatting, clear the rest
        if paragraph.runs:
            fmt = paragraph.runs[0].font
            font_name = fmt.name
            font_size = fmt.size
            font_bold = fmt.bold
            # Clear all runs and set consolidated text
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = full_text
            # Re-apply formatting
            paragraph.runs[0].font.name = font_name
            if font_size:
                paragraph.runs[0].font.size = font_size
            if font_bold is not None:
                paragraph.runs[0].font.bold = font_bold
        else:
            paragraph.text = full_text

    return replaced


def _replace_in_table(table, replacements: dict[str, str]):
    """Replace placeholders in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)


def _replace_placeholders(doc: DocxDocument, replacements: dict[str, str]):
    """Replace all {{PLACEHOLDER}} markers throughout the document."""
    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Tables
    for table in doc.tables:
        _replace_in_table(table, replacements)

    # Headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, replacements)
        for table in section.header.tables:
            _replace_in_table(table, replacements)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, replacements)
        for table in section.footer.tables:
            _replace_in_table(table, replacements)


# ──────────────────────────────────────────────────────────────
# Section injection — find placeholder paragraph and add content after it
# ──────────────────────────────────────────────────────────────

def _find_and_inject_section(doc: DocxDocument, placeholder: str, content: str):
    """
    Find a paragraph containing {{placeholder}} and replace it with
    the section content (multiple paragraphs).
    """
    marker = "{{" + placeholder + "}}"
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text:
            # Clear the placeholder paragraph and use it for first line
            lines = [l for l in content.split("\n") if l.strip()] if content else []
            if lines:
                para.text = lines[0]
                # Add remaining lines after this paragraph
                for line in lines[1:]:
                    new_para = doc.add_paragraph(line)
                    new_para.style = para.style
                    # Move the new paragraph right after the current one
                    para._element.addnext(new_para._element)
                    para = new_para
            else:
                para.text = ""
            return True
    return False


# ──────────────────────────────────────────────────────────────
# Populate revision history table
# ──────────────────────────────────────────────────────────────

def _populate_revision_history(doc: DocxDocument, history_entries: list[dict]):
    """
    Find the revision history table and populate it.
    Each entry: {"revision": "01", "date": "23/02/2026", "changes": "...", "responsible": "..."}
    """
    if not history_entries:
        return

    # Find a table that looks like revision history
    # (has columns: Revisão, Data, Alterações, Responsável)
    for table in doc.tables:
        header_text = " ".join(
            cell.text.strip().lower() for cell in table.rows[0].cells
        ) if table.rows else ""
        if "revis" in header_text and ("altera" in header_text or "data" in header_text):
            # Found the history table — add rows
            for entry in history_entries:
                row = table.add_row()
                cells = row.cells
                if len(cells) >= 4:
                    cells[0].text = str(entry.get("revision", ""))
                    cells[1].text = str(entry.get("date", ""))
                    cells[2].text = str(entry.get("changes", ""))
                    cells[3].text = str(entry.get("responsible", ""))
                elif len(cells) >= 3:
                    cells[0].text = str(entry.get("revision", ""))
                    cells[1].text = str(entry.get("date", ""))
                    cells[2].text = str(entry.get("changes", ""))
            return


# ──────────────────────────────────────────────────────────────
# Populate approval/consensus table
# ──────────────────────────────────────────────────────────────

def _populate_approval_table(doc: DocxDocument, approvers: list[dict]):
    """
    Find the approval table and populate it.
    Each approver: {"type": "A", "date": "23/02/2026", "name": "...", "sector": "...", "signature": ""}
    """
    if not approvers:
        return

    for table in doc.tables:
        header_text = " ".join(
            cell.text.strip().lower() for cell in table.rows[0].cells
        ) if table.rows else ""
        if "aprova" in header_text or "consenso" in header_text or "assinatura" in header_text:
            for approver in approvers:
                row = table.add_row()
                cells = row.cells
                if len(cells) >= 5:
                    cells[0].text = str(approver.get("type", "A"))
                    cells[1].text = str(approver.get("date", ""))
                    cells[2].text = str(approver.get("name", ""))
                    cells[3].text = str(approver.get("sector", ""))
                    cells[4].text = str(approver.get("signature", ""))
                elif len(cells) >= 3:
                    cells[0].text = str(approver.get("name", ""))
                    cells[1].text = str(approver.get("sector", ""))
                    cells[2].text = str(approver.get("date", ""))
            return


# ──────────────────────────────────────────────────────────────
# ODT to DOCX conversion
# ──────────────────────────────────────────────────────────────

def convert_odt_to_docx(odt_path: str, output_dir: str) -> str:
    """Convert .odt to .docx using LibreOffice headless."""
    os.makedirs(output_dir, exist_ok=True)
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", output_dir, odt_path],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

    base_name = os.path.splitext(os.path.basename(odt_path))[0]
    return os.path.join(output_dir, f"{base_name}.docx")


def convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
    """Convert .docx to .pdf using LibreOffice headless."""
    os.makedirs(output_dir, exist_ok=True)
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice PDF conversion failed: {result.stderr}")

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    return os.path.join(output_dir, f"{base_name}.pdf")


# ──────────────────────────────────────────────────────────────
# Find placeholders in a template
# ──────────────────────────────────────────────────────────────

def find_placeholders(template_path: str) -> list[str]:
    """Scan a .docx template and return all {{PLACEHOLDER}} names found."""
    doc = DocxDocument(template_path)
    found = set()

    def _scan_text(text: str):
        for m in PLACEHOLDER_RE.finditer(text):
            found.add(m.group(1))

    for para in doc.paragraphs:
        _scan_text(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_text(para.text)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _scan_text(para.text)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _scan_text(para.text)
        for para in section.footer.paragraphs:
            _scan_text(para.text)

    return sorted(found)


# ──────────────────────────────────────────────────────────────
# Main formatting pipeline
# ──────────────────────────────────────────────────────────────

def format_with_template(
    template_path: str,
    sections: dict[str, str],
    metadata: dict,
    history_entries: list[dict] | None = None,
    approvers: list[dict] | None = None,
    source_images: list[dict] | None = None,
    output_path: str = "",
) -> str:
    """
    Full formatting pipeline:
    1. Load template .docx (convert .odt if needed)
    2. Replace header placeholders (TITULO, CODIGO, REVISAO, DATA, SETOR)
    3. Inject section content at {{SECTION}} markers
    4. Populate revision history table
    5. Populate approval table
    6. Save formatted .docx

    Returns the path to the formatted .docx file.
    """
    temp_dir = os.path.join(settings.STORAGE_PATH, "temp")
    os.makedirs(temp_dir, exist_ok=True)

    # Convert .odt if needed
    working_template = template_path
    if template_path.lower().endswith(".odt"):
        working_template = convert_odt_to_docx(template_path, temp_dir)

    # Open template
    doc = DocxDocument(working_template)

    # 1. Header placeholders
    header_replacements = {
        "TITULO": metadata.get("title", ""),
        "CODIGO": metadata.get("code", ""),
        "REVISAO": metadata.get("revision", ""),
        "DATA": metadata.get("date", ""),
        "SETOR": metadata.get("sector", ""),
        "DATA_VIGOR": metadata.get("date", ""),
    }
    _replace_placeholders(doc, header_replacements)

    # 2. Section content — inject via placeholder replacement
    section_replacements = {}
    for key, content in sections.items():
        section_key = key.upper().replace(" ", "_")
        section_replacements[section_key] = content

    # Try both paragraph injection and simple replacement
    for key, content in section_replacements.items():
        if not _find_and_inject_section(doc, key, content):
            # Fallback: try simple text replacement
            _replace_placeholders(doc, {key: content})

    # 3. Revision history
    if history_entries:
        _populate_revision_history(doc, history_entries)

    # 4. Approval table
    if approvers:
        _populate_approval_table(doc, approvers)

    # 5. Clean remaining placeholders (replace with empty)
    remaining = {}
    for para in doc.paragraphs:
        for m in PLACEHOLDER_RE.finditer(para.text):
            remaining[m.group(1)] = ""
    if remaining:
        _replace_placeholders(doc, remaining)

    # Save
    os.makedirs(os.path.dirname(output_path) if output_path else temp_dir, exist_ok=True)
    if not output_path:
        output_path = os.path.join(temp_dir, "formatted_output.docx")
    doc.save(output_path)

    return output_path


async def format_document_with_template(
    template_path: str,
    structured_content: dict,
    metadata: dict,
    source_docx_path: str | None,
    changelog_entries: list[dict] | None,
    approval_data: list[dict] | None,
    output_docx_path: str,
    output_pdf_path: str,
) -> tuple[str, str]:
    """
    High-level async wrapper for the full format pipeline.
    Returns (docx_path, pdf_path).
    """
    # Extract sections from structured content
    sections = {}
    key_mapping = {
        "OBJETIVO_E_ABRANGENCIA": "OBJETIVO",
        "DOCUMENTOS_COMPLEMENTARES": "DOCUMENTOS_COMPLEMENTARES",
        "DEFINICOES": "DEFINICOES",
        "DESCRICAO_DAS_ATIVIDADES": "ATIVIDADES",
        "RESPONSABILIDADES": "RESPONSABILIDADES",
        "CARACTERISTICAS": "CARACTERISTICAS",
        "CONDICOES_DE_SEGURANCA": "SEGURANCA",
        "CONDICOES_DE_ARMAZENAMENTO": "ARMAZENAMENTO",
        "ALTERACOES": "ALTERACOES",
    }
    for sec in structured_content.get("sections", []):
        title = sec.get("title", "Content")
        content = sec.get("content", "")
        key = _strip_accents(title.upper().replace(" ", "_"))
        mapped_key = key_mapping.get(key, key)
        sections[mapped_key] = content

    # Build history entries from changelog
    history = []
    if changelog_entries:
        for entry in changelog_entries:
            history.append({
                "revision": entry.get("revision", ""),
                "date": entry.get("date", ""),
                "changes": entry.get("changes", ""),
                "responsible": entry.get("responsible", ""),
            })

    # Extract images from source document
    images = None
    if source_docx_path and os.path.exists(source_docx_path):
        images = extract_images_from_docx(source_docx_path)

    # Run formatting
    docx_path = format_with_template(
        template_path=template_path,
        sections=sections,
        metadata=metadata,
        history_entries=history,
        approvers=approval_data,
        source_images=images,
        output_path=output_docx_path,
    )

    # Convert to PDF
    pdf_path = ""
    try:
        output_dir = os.path.dirname(output_pdf_path)
        pdf_path = convert_docx_to_pdf(docx_path, output_dir)
        # Rename if needed
        if pdf_path != output_pdf_path:
            os.rename(pdf_path, output_pdf_path)
            pdf_path = output_pdf_path
    except Exception:
        pdf_path = ""

    return docx_path, pdf_path
